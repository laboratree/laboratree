"""Multi-format extraction — pull tables and text out of CSV, Excel, DOCX, and PDF bytes.

Pure functions (no I/O beyond the provided bytes) so they are trivially testable and reusable
by both the Signal API and the Paper Lab's ingestion.
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field
from pathlib import Path

import pandas as pd


@dataclass
class ExtractedTable:
    name: str          # human label, e.g. "sales:Sheet1" or "report:table1"
    df: pd.DataFrame
    source: str        # originating filename
    kind: str          # csv | sheet | table


@dataclass
class ExtractResult:
    source: str
    tables: list[ExtractedTable] = field(default_factory=list)
    texts: list[str] = field(default_factory=list)  # free-text blocks (docx/pdf)


def _stem(name: str) -> str:
    return Path(name).stem or "file"


def extract_csv(name: str, data: bytes) -> ExtractResult:
    for enc in ("utf-8", "latin-1"):
        try:
            df = pd.read_csv(io.BytesIO(data), encoding=enc)
            break
        except (UnicodeDecodeError, pd.errors.ParserError):
            continue
    else:
        df = pd.read_csv(io.BytesIO(data), encoding="utf-8", engine="python", on_bad_lines="skip")
    return ExtractResult(source=name, tables=[ExtractedTable(_stem(name), df, name, "csv")])


def extract_excel(name: str, data: bytes) -> ExtractResult:
    sheets = pd.read_excel(io.BytesIO(data), sheet_name=None)
    tables = [
        ExtractedTable(f"{_stem(name)}:{sheet}", df, name, "sheet")
        for sheet, df in sheets.items()
    ]
    return ExtractResult(source=name, tables=tables)


def extract_docx(name: str, data: bytes) -> ExtractResult:
    import docx

    doc = docx.Document(io.BytesIO(data))
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    tables: list[ExtractedTable] = []
    for i, table in enumerate(doc.tables, start=1):
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        header, *body = rows
        df = pd.DataFrame(body, columns=_dedupe(header)) if body else pd.DataFrame(columns=header)
        tables.append(ExtractedTable(f"{_stem(name)}:table{i}", df, name, "table"))
    return ExtractResult(source=name, tables=tables, texts=texts)


def extract_pdf(name: str, data: bytes) -> ExtractResult:
    import pdfplumber

    tables: list[ExtractedTable] = []
    texts: list[str] = []
    counter = 1
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            txt = page.extract_text() or ""
            if txt.strip():
                texts.append(txt)
            for raw in page.extract_tables():
                if not raw:
                    continue
                header, *body = raw
                header = _dedupe([(h or f"col{j}") for j, h in enumerate(header)])
                df = pd.DataFrame(body, columns=header) if body else pd.DataFrame(columns=header)
                tables.append(ExtractedTable(f"{_stem(name)}:table{counter}", df, name, "table"))
                counter += 1
    # Scanned / image-only PDF: no text layer -> fall back to OCR if an engine is available.
    if not texts:
        try:
            from ...ocr import ocr_available, ocr_pdf_pages

            if ocr_available():
                texts = ocr_pdf_pages(data)
        except Exception:
            pass
    return ExtractResult(source=name, tables=tables, texts=texts)


def extract_image(name: str, data: bytes) -> ExtractResult:
    from ...ocr import ocr_available, ocr_image_bytes

    if not ocr_available():
        raise ValueError(
            "image OCR requires the tesseract binary (not installed); "
            "see laboratree/ocr/README.md"
        )
    text = ocr_image_bytes(data)
    return ExtractResult(source=name, texts=[text] if text.strip() else [])


_EXTRACTORS = {
    ".csv": extract_csv,
    ".tsv": extract_csv,
    ".xlsx": extract_excel,
    ".xls": extract_excel,
    ".xlsm": extract_excel,
    ".docx": extract_docx,
    ".pdf": extract_pdf,
    ".png": extract_image,
    ".jpg": extract_image,
    ".jpeg": extract_image,
    ".tif": extract_image,
    ".tiff": extract_image,
    ".bmp": extract_image,
}


def supported_suffixes() -> list[str]:
    return sorted(_EXTRACTORS)


def extract_file(name: str, data: bytes) -> ExtractResult:
    suffix = Path(name).suffix.lower()
    extractor = _EXTRACTORS.get(suffix)
    if extractor is None:
        raise ValueError(f"unsupported file type: {suffix or name!r}")
    return extractor(name, data)


def _dedupe(cols: list) -> list[str]:
    seen: dict[str, int] = {}
    out: list[str] = []
    for c in cols:
        c = str(c).strip() or "col"
        if c in seen:
            seen[c] += 1
            out.append(f"{c}.{seen[c]}")
        else:
            seen[c] = 0
            out.append(c)
    return out
