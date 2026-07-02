"""Phase 6 tests: ML models, data-fetch agent (+HITL), walkthrough, and the experiment flow."""

from __future__ import annotations

import io
import uuid

import docx
import pandas as pd
from fastapi.testclient import TestClient

from laboratree.core.registry import REGISTRY
from laboratree.labs.paper.experiment.fetch import (
    DataFetchAgent,
    DatasetRef,
    SklearnToyResolver,
    extract_dataset_refs,
)
from laboratree.labs.paper.experiment.walkthrough import build_walkthrough, default_walkthrough
from laboratree.main import app
from laboratree_sdk import RunContext


class _Sink:
    def __init__(self):
        self.records = []

    def record(self, *, label, value, kind="metric", **meta):
        self.records.append((label, value))
        return f"e{len(self.records)}"


# ---------------- ML models ----------------

def test_logistic_regression_component_runs():
    from sklearn.datasets import load_iris

    df = load_iris(as_frame=True).frame
    ctx = RunContext(run_id="r", org_id="o", params={"target": "target"},
                     inputs={"dataset": df}, evidence=_Sink())
    out = REGISTRY.create("model.ml.logistic_regression").run(ctx)
    assert "accuracy" in out["metrics"]
    assert 0.0 <= out["metrics"]["accuracy"] <= 1.0


def test_linear_regression_component_runs():
    df = pd.DataFrame({"x": range(50)})
    df["y"] = 2 * df["x"] + 1
    ctx = RunContext(run_id="r", org_id="o", params={"target": "y"},
                     inputs={"dataset": df}, evidence=_Sink())
    out = REGISTRY.create("model.ml.linear_regression").run(ctx)
    assert out["metrics"]["r2"] > 0.9


# ---------------- fetch agent ----------------

def test_sklearn_resolver_fetches_iris():
    fr = SklearnToyResolver().try_fetch(DatasetRef("iris"))
    assert fr is not None and fr.filename == "iris.csv"


def test_agent_resolves_known_and_flags_unknown_for_hitl():
    outcome = DataFetchAgent().resolve([DatasetRef("iris"), DatasetRef("secret_survey_2021")])
    assert len(outcome.fetched) == 1
    assert outcome.needs_human
    assert "download" in outcome.unresolved[0].instructions.lower()


def test_extract_refs_finds_data_url():
    refs = extract_dataset_refs("Data at https://example.com/data.csv here.", None)
    assert any(r.url and r.url.endswith("data.csv") for r in refs)


# ---------------- walkthrough ----------------

def test_default_walkthrough_maps_model_component():
    card = {"models_used": ["Logistic regression"], "target_variable": "y",
            "data_sources": ["iris"], "preprocessing": ["scale"], "independent_variables": ["a"]}
    steps = default_walkthrough(card)
    kinds = [s["kind"] for s in steps]
    assert kinds[0] == "data" and "model" in kinds and kinds[-1] == "inference"
    model = next(s for s in steps if s["kind"] == "model")
    assert model["component_id"] == "model.ml.logistic_regression"


def test_build_walkthrough_falls_back_on_bad_llm():
    card = {"models_used": ["OLS"], "target_variable": "y"}
    steps = build_walkthrough(card, complete_fn=lambda s, p: "not json")
    assert steps == default_walkthrough(card)


# ---------------- experiment API flow ----------------

def _paper_docx() -> bytes:
    d = docx.Document()
    d.add_paragraph("Churn prediction with logistic regression on the iris-like dataset.")
    for _ in range(4):
        d.add_paragraph("We use logistic regression. Target is target. Data source: iris. AUC 0.84.")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _fake_complete(system: str, prompt: str, **kw) -> str:
    if "dataset references" in system:
        return '[{"name": "iris"}, {"name": "secret_survey_2021"}]'
    if "reconstruct" in system:
        return "[]"  # force fallback to the card-derived default walkthrough
    if "Paper Card" in system:
        return ('{"problem_statement": "Predict target", "models_used": ["Logistic regression"], '
                '"target_variable": "target", "results": "AUC 0.84", "data_sources": ["iris"], '
                '"preprocessing": ["scale"], "independent_variables": ["sepal length (cm)"]}')
    return "ok"


def _register(client):
    email = f"user-{uuid.uuid4().hex[:10]}@example.com"
    r = client.post("/api/auth/register",
                    json={"email": email, "password": "supersecret1", "full_name": "E"})
    h = {"Authorization": f"Bearer {r.json()['access_token']}"}
    p = client.post("/api/projects", json={"name": "Exp"}, headers=h)
    return h, p.json()["id"]


def test_experiment_fetch_hitl_and_node_run(monkeypatch):
    from laboratree.labs.paper import llm as paper_llm

    monkeypatch.setattr(paper_llm, "default_complete", _fake_complete)
    monkeypatch.setattr(paper_llm, "default_embed", lambda texts: (_ for _ in ()).throw(RuntimeError()))

    with TestClient(app) as client:
        h, project_id = _register(client)
        up = client.post(f"/api/projects/{project_id}/papers", headers=h,
                         files={"file": ("p.docx", _paper_docx(),
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
        pid = up.json()["id"]
        client.post(f"/api/papers/{pid}/card", headers=h)  # generate card (fake LLM)

        # start experiment: iris auto-fetched, secret_survey_2021 needs a human
        exp = client.post(f"/api/papers/{pid}/experiment", headers=h).json()
        assert exp["status"] == "awaiting_data"
        assert exp["gate_id"]
        fetched_names = {f["name"] for f in exp["fetch_report"]["fetched"]}
        assert "iris" in fetched_names
        assert any(u["name"] == "secret_survey_2021" for u in exp["fetch_report"]["unresolved"])
        exp_id = exp["id"]

        # a model node maps to a runnable component
        model_node = next(n for n in exp["walkthrough"] if n["kind"] == "model" and n.get("component_id"))
        iris_ds = next(f["dataset_id"] for f in exp["fetch_report"]["fetched"] if f["name"] == "iris")

        # run the model node on the fetched iris data -> Evidence-locked metrics + compare
        run = client.post(
            f"/api/experiments/{exp_id}/nodes/{model_node['id']}/run",
            headers=h, json={"dataset_id": iris_ds},
        )
        assert run.status_code == 201, run.text
        assert "accuracy" in run.json()["metrics"]
        assert run.json()["paper_reported"] == "AUC 0.84"

        # HITL: upload the missing dataset -> experiment becomes ready
        csv = pd.DataFrame({"a": [1, 2], "b": [3, 4]}).to_csv(index=False).encode()
        done = client.post(
            f"/api/experiments/{exp_id}/data?name=secret_survey_2021",
            headers=h, files={"file": ("survey.csv", csv, "text/csv")},
        )
        assert done.status_code == 201, done.text
        assert done.json()["status"] == "ready"
        assert done.json()["fetch_report"]["unresolved"] == []
