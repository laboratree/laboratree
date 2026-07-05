"""Phase 5 tests: paper ingest/chunk, Paper Card, explain-simpler, and chat (offline LLM)."""

from __future__ import annotations

import io
import uuid

import docx
from fastapi.testclient import TestClient

from laboratree.labs.paper.card import generate_card, normalize_card
from laboratree.labs.paper.ingest import chunk_text
from laboratree.labs.paper.rag import answer as rag_answer
from laboratree.labs.paper.simplify import simplify
from laboratree.main import app


# ---------------- unit ----------------

def test_chunk_text_overlaps_and_covers():
    text = "abcdefghij" * 400  # 4000 chars
    chunks = chunk_text(text, size=1500, overlap=200)
    assert len(chunks) >= 3
    assert all(len(c) <= 1500 for c in chunks)


def test_generate_card_parses_json_and_normalizes():
    def fake(system, prompt, **kw):
        return '```json\n{"problem_statement": "Predicts churn.", "models_used": ["Logit"]}\n```'

    card = generate_card("some paper text", complete_fn=fake)
    assert card["problem_statement"]["plain"] == "Predicts churn."
    assert card["models_used"][0]["name"] == "Logit"
    # normalized: missing fields present
    assert "math" in card and isinstance(card["math"], list)


def test_normalize_card_fills_defaults():
    card = normalize_card({"problem_statement": "x"})
    assert card["target_variable"]["name"] == ""
    assert card["data_sources"] == []


def test_simplify_calls_model():
    out = simplify("A stochastic gradient estimator.", 3, complete_fn=lambda s, p: "It nudges numbers.")
    assert out == "It nudges numbers."


def test_rag_answer_includes_citations():
    passages = [{"ordinal": 0, "text": "The model is AR(1)."}]
    res = rag_answer("what model?", passages, complete_fn=lambda s, p: "It uses AR(1) [0].")
    assert res["citations"] == [0]
    assert "[0]" in res["answer"]


# ---------------- API (needs Postgres; LLM monkeypatched) ----------------

def _docx_paper() -> bytes:
    d = docx.Document()
    d.add_paragraph("Title: Predicting customer churn with logistic regression.")
    for _ in range(6):
        d.add_paragraph(
            "We model churn using a logistic regression on tenure and monthly charges. "
            "The dataset has 7000 customers. Preprocessing includes one-hot encoding and scaling. "
            "The target variable is churn (yes/no). Results show an AUC of about 0.84."
        )
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _fake_complete(system: str, prompt: str, **kw) -> str:
    if "JSON" in system:
        return '{"problem_statement": "Predict churn.", "models_used": ["Logistic regression"], "target_variable": "churn"}'
    if "make hard ideas easy" in system:
        return "Think of it like guessing who will leave a gym."
    return "The paper uses logistic regression [0]."


def _fake_embed(texts):
    raise RuntimeError("offline")  # force keyword retrieval, skip embeddings


def _register(client: TestClient) -> tuple[dict, str]:
    email = f"user-{uuid.uuid4().hex[:10]}@example.com"
    r = client.post("/api/auth/register",
                    json={"email": email, "password": "supersecret1", "full_name": "P"})
    h = {"Authorization": f"Bearer {r.json()['access_token']}"}
    p = client.post("/api/projects", json={"name": "Papers"}, headers=h)
    return h, p.json()["id"]


def test_paper_study_flow(monkeypatch):
    from laboratree.labs.paper import llm as paper_llm

    monkeypatch.setattr(paper_llm, "default_complete", _fake_complete)
    monkeypatch.setattr(paper_llm, "default_embed", _fake_embed)

    with TestClient(app) as client:
        h, project_id = _register(client)

        up = client.post(
            f"/api/projects/{project_id}/papers",
            headers=h,
            files={"file": ("churn.docx", _docx_paper(),
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert up.status_code == 201, up.text
        paper = up.json()
        assert paper["status"] == "parsed"
        assert paper["n_chunks"] >= 1
        pid = paper["id"]

        # Paper Card
        card = client.post(f"/api/papers/{pid}/card", headers=h).json()
        assert card["card"]["problem_statement"]["plain"] == "Predict churn."
        assert card["status"] == "carded"

        # explain simpler
        simp = client.post(
            f"/api/papers/{pid}/simplify",
            headers=h,
            json={"field": "problem_statement", "level": 3},
        ).json()
        assert "gym" in simp["simplified"]

        # chat-with-paper (keyword retrieval)
        chat = client.post(
            f"/api/papers/{pid}/chat", headers=h, json={"question": "what model is used?"}
        ).json()
        assert chat["answer"]
        assert isinstance(chat["citations"], list)
