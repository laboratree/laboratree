"""Phase 4 tests: multi-format extraction, consolidation, Signal API + artifact download."""

from __future__ import annotations

import io
import uuid

import docx
import pandas as pd
from fastapi.testclient import TestClient

from laboratree.core.registry import REGISTRY
from laboratree.core.storage import get_blob_store
from laboratree.labs.signal.consolidate import consolidate
from laboratree.labs.signal.extract import extract_docx, extract_file
from laboratree.main import app
from laboratree_sdk import RunContext


# ---------- fixtures (in-memory files) ----------

def _csv_bytes() -> bytes:
    return pd.DataFrame({"a": [1, 2, 3], "b": ["x", "y", "z"]}).to_csv(index=False).encode()


def _xlsx_bytes() -> bytes:
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as w:
        pd.DataFrame({"m": [1, 2]}).to_excel(w, index=False, sheet_name="Metrics")
    return buf.getvalue()


def _docx_bytes() -> bytes:
    d = docx.Document()
    d.add_paragraph("Executive summary text.")
    t = d.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text, t.rows[0].cells[1].text = "col1", "col2"
    t.rows[1].cells[0].text, t.rows[1].cells[1].text = "10", "20"
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


class _ListSink:
    def __init__(self):
        self.records = []

    def record(self, *, label, value, kind="metric", **meta):
        self.records.append({"label": label, "value": value})
        return f"ev{len(self.records)}"


# ---------- extraction + consolidation (pure) ----------

def test_extract_docx_tables_and_text():
    res = extract_docx("report.docx", _docx_bytes())
    assert res.texts and "Executive summary" in res.texts[0]
    assert res.tables and list(res.tables[0].df.columns) == ["col1", "col2"]


def test_consolidate_builds_master_workbook_with_dictionary():
    files = [("sales.csv", _csv_bytes()), ("kpis.xlsx", _xlsx_bytes()), ("memo.docx", _docx_bytes())]
    result = consolidate(files)
    assert result.n_tables >= 3
    assert result.texts >= 1
    sheets = pd.read_excel(io.BytesIO(result.workbook), sheet_name=None)
    assert "Data Dictionary" in sheets
    assert "Text Blocks" in sheets
    dictionary = sheets["Data Dictionary"]
    assert set(dictionary["source"]) >= {"sales.csv", "kpis.xlsx"}


def test_unsupported_file_is_reported_not_raised():
    result = consolidate([("notes.rtf", b"junk")])
    assert result.errors and result.errors[0]["source"] == "notes.rtf"


# ---------- file connector component ----------

def test_file_connector_reads_blob():
    key = f"tests/{uuid.uuid4().hex}.csv"
    get_blob_store().put(key, _csv_bytes())
    ctx = RunContext(
        run_id="r", org_id="o",
        params={"storage_key": key, "filename": "data.csv"},
        blobs=get_blob_store(), evidence=_ListSink(),
    )
    out = REGISTRY.create("connector.file").run(ctx)
    assert list(out["dataset"].columns) == ["a", "b"]


# ---------- Signal API (needs Postgres) ----------

def _register(client: TestClient) -> tuple[dict, str]:
    email = f"user-{uuid.uuid4().hex[:10]}@example.com"
    r = client.post("/api/auth/register",
                    json={"email": email, "password": "supersecret1", "full_name": "S"})
    token = r.json()["access_token"]
    h = {"Authorization": f"Bearer {token}"}
    p = client.post("/api/projects", json={"name": "Signal"}, headers=h)
    return h, p.json()["id"]


def test_signal_consolidate_endpoint_and_download():
    with TestClient(app) as client:
        h, project_id = _register(client)
        files = [
            ("files", ("sales.csv", _csv_bytes(), "text/csv")),
            ("files", ("kpis.xlsx", _xlsx_bytes(),
                       "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")),
        ]
        r = client.post(f"/api/projects/{project_id}/signal/consolidate", files=files, headers=h)
        assert r.status_code == 201, r.text
        body = r.json()
        assert body["summary"]["n_tables"] >= 2

        ev = client.get(f"/api/runs/{body['run_id']}/evidence", headers=h).json()
        labels = {e["label"]: e["value"] for e in ev}
        assert labels.get("consolidated_tables", 0) >= 2

        dl = client.get(body["download_url"], headers=h)
        assert dl.status_code == 200
        sheets = pd.read_excel(io.BytesIO(dl.content), sheet_name=None)
        assert "Data Dictionary" in sheets
