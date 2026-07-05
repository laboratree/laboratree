"""Generate a richer set of sample files for testing every Lab.

Run (fpdf2 is pulled in just for this script, not added to the project):
  cd apps/api
  uv run --with fpdf2 python ../../docs/samples/make_sample_files.py

Creates in docs/samples/:
  churn_clean.csv          classification, no leakage      (Model/Red-Team/Pipeline/Insight)
  housing.csv              regression                      (Model linear regression)
  campaign_timeseries.csv  weekly sales w/ an intervention (Trend causal impact, index 10)
  survey_responses.csv     survey data                     (Insight/Decision)
  messy_finance.xlsx       2 sheets (revenue, costs)       (Signal consolidation)
  quarterly_report.docx    text + a table                  (Signal consolidation)
  sample_paper_wine.pdf    a short paper (wine dataset)    (Paper Lab; auto-fetch 'wine')
"""

from pathlib import Path

import numpy as np
import pandas as pd

OUT = Path(__file__).parent
rng = np.random.default_rng(7)


def _sigmoid(x):
    return 1 / (1 + np.exp(-x))


# 1) classification, no leakage
n = 40
tenure = rng.integers(1, 72, n)
monthly = rng.uniform(20, 120, n).round(2)
support = rng.integers(0, 8, n)
senior = rng.integers(0, 2, n)
logit = -2.5 + 0.03 * monthly + 0.35 * support - 0.04 * tenure + 0.5 * senior
churned = (rng.uniform(0, 1, n) < _sigmoid(logit)).astype(int)
pd.DataFrame({
    "tenure_months": tenure, "monthly_charges": monthly, "support_calls": support,
    "is_senior": senior, "churned": churned,
}).to_csv(OUT / "churn_clean.csv", index=False)

# 2) regression
size = rng.integers(600, 3000, n)
beds = rng.integers(1, 6, n)
age = rng.integers(0, 60, n)
price = (120 * size + 15000 * beds - 800 * age + rng.normal(0, 15000, n)).round(0).astype(int)
pd.DataFrame({"size_sqft": size, "bedrooms": beds, "age_years": age, "price": price}).to_csv(
    OUT / "housing.csv", index=False)

# 3) time series with an intervention at week 11 (index 10)
weeks = np.arange(1, 21)
sales = (200 + 3 * weeks + rng.normal(0, 6, 20)).round(0)
sales[10:] += 45  # campaign lift
pd.DataFrame({"week": weeks, "sales": sales.astype(int)}).to_csv(
    OUT / "campaign_timeseries.csv", index=False)

# 4) survey responses
groups = rng.choice(["18-29", "30-44", "45-59", "60+"], n)
sat = rng.integers(1, 6, n)
rec = (sat >= 4).astype(int)
spend = (sat * 12 + rng.normal(0, 8, n)).round(2)
pd.DataFrame({
    "respondent_id": np.arange(1, n + 1), "satisfaction": sat, "would_recommend": rec,
    "monthly_spend": spend, "age_group": groups,
}).to_csv(OUT / "survey_responses.csv", index=False)

# 5) messy multi-sheet Excel (Signal)
with pd.ExcelWriter(OUT / "messy_finance.xlsx", engine="openpyxl") as w:
    pd.DataFrame({"quarter": ["Q1", "Q2", "Q3", "Q4"], "revenue": [120, 135, 150, 172]}).to_excel(
        w, sheet_name="revenue", index=False)
    pd.DataFrame({"quarter": ["Q1", "Q2", "Q3", "Q4"], "cost": [80, 88, 96, 101]}).to_excel(
        w, sheet_name="costs", index=False)

# 6) Word doc with a table (Signal)
import docx

d = docx.Document()
d.add_heading("Quarterly Report", level=1)
d.add_paragraph("Summary of headcount by department for the last quarter.")
t = d.add_table(rows=1, cols=2)
t.rows[0].cells[0].text, t.rows[0].cells[1].text = "department", "headcount"
for dept, hc in [("Engineering", "24"), ("Sales", "13"), ("Research", "9")]:
    row = t.add_row().cells
    row[0].text, row[1].text = dept, hc
d.save(OUT / "quarterly_report.docx")

# 7) a real text-layer PDF paper (mentions the 'wine' dataset -> auto-fetchable)
from fpdf import FPDF

pdf = FPDF()
pdf.add_page()
pdf.set_font("Helvetica", "B", 15)
pdf.multi_cell(0, 9, "Predicting Wine Cultivar with Logistic Regression")
pdf.ln(2)
pdf.set_font("Helvetica", size=11)
for para in [
    "Abstract. We classify wines by cultivar using logistic regression on the wine dataset.",
    "Data. The wine dataset has 178 samples and 13 numeric chemical features (alcohol, "
    "malic acid, ash, magnesium, flavanoids, color intensity, proline, and others). The target "
    "variable is the cultivar class. Features are standardized before modeling.",
    "Method. We fit a multinomial logistic regression with a 75/25 train/test split and report "
    "accuracy on the held-out set. The independent variables are the 13 measurements.",
    "Results. The model reaches about 0.97 accuracy. Flavanoids and proline are the most "
    "informative features. We conclude logistic regression is a strong baseline for this task.",
]:
    pdf.multi_cell(0, 6, para)
    pdf.ln(1)
pdf.output(str(OUT / "sample_paper_wine.pdf"))

print("wrote sample files to", OUT)
for p in sorted(OUT.glob("*")):
    print(" ", p.name)
