"""Generate a small sample research paper (DOCX) for testing the Paper Lab.

Run:  cd apps/api && uv run python ../../docs/samples/make_sample_paper.py
Creates docs/samples/sample_paper.docx — mentions the 'iris' dataset + logistic regression so
the Paper Experiment auto-fetch agent has something to retrieve.
"""

from pathlib import Path

import docx

out = Path(__file__).parent / "sample_paper.docx"
d = docx.Document()
d.add_heading("Classifying Iris Flowers with Logistic Regression", level=1)
d.add_paragraph(
    "Abstract. We study multiclass classification of iris flowers using logistic regression on the "
    "classic iris dataset. The target variable is the species (setosa, versicolor, virginica)."
)
d.add_heading("Data", level=2)
d.add_paragraph(
    "We use the iris dataset (150 samples, 4 numeric features: sepal length, sepal width, petal "
    "length, petal width). Preprocessing includes standardization of the features."
)
d.add_heading("Method", level=2)
d.add_paragraph(
    "We fit a multinomial logistic regression. The independent variables are the four measurements; "
    "the target variable is the species class. We evaluate with a 75/25 train/test split."
)
d.add_heading("Results", level=2)
d.add_paragraph(
    "The model achieves about 0.95 accuracy on the held-out test set. Petal length and petal width "
    "are the most informative features."
)
d.save(out)
print("wrote", out)
